"""
Модуль для работы с договорами: сохранение, экспорт, история
"""

import json
import os
from datetime import datetime
from typing import Optional, Dict, List
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ContractManager:
    """Менеджер для работы с договорами"""
    
    def __init__(self, storage_dir: str = "contracts"):
        self.storage_dir = storage_dir
        self.metadata_file = os.path.join(storage_dir, "metadata.json")
        
        # Создаем директорию если её нет
        Path(storage_dir).mkdir(exist_ok=True)
        
        # Инициализируем метаданные
        if not os.path.exists(self.metadata_file):
            self._save_metadata({})
    
    def _load_metadata(self) -> Dict:
        """Загружает метаданные договоров"""
        if os.path.exists(self.metadata_file):
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _save_metadata(self, data: Dict):
        """Сохраняет метаданные договоров"""
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def save_contract(self, 
                     contract_text: str,
                     client_data: str,
                     contract_type: str,
                     user_id: int,
                     username: str) -> str:
        """
        Сохраняет договор и возвращает ID
        
        Args:
            contract_text: Текст договора
            client_data: Данные клиента
            contract_type: Тип договора
            user_id: ID пользователя Telegram
            username: Имя пользователя
            
        Returns:
            contract_id: ID договора
        """
        contract_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Сохраняем текст договора
        filename = f"contract_{contract_id}.txt"
        filepath = os.path.join(self.storage_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(contract_text)
        
        # Обновляем метаданные
        metadata = self._load_metadata()
        metadata[contract_id] = {
            'filename': filename,
            'contract_type': contract_type,
            'user_id': user_id,
            'username': username,
            'client_data': client_data,
            'created_at': datetime.now().isoformat(),
            'formats': ['txt']
        }
        self._save_metadata(metadata)
        
        return contract_id
    
    def get_contract(self, contract_id: str) -> Optional[str]:
        """Получает текст договора по ID"""
        metadata = self._load_metadata()
        if contract_id not in metadata:
            return None
        
        filepath = os.path.join(self.storage_dir, metadata[contract_id]['filename'])
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        return None
    
    def get_user_contracts(self, user_id: int) -> List[Dict]:
        """Получает все договоры пользователя"""
        metadata = self._load_metadata()
        user_contracts = []
        
        for contract_id, info in metadata.items():
            if info['user_id'] == user_id:
                user_contracts.append({
                    'id': contract_id,
                    'type': info['contract_type'],
                    'created_at': info['created_at'],
                    'client': info['client_data']
                })
        
        return sorted(user_contracts, key=lambda x: x['created_at'], reverse=True)
    
    def export_to_word(self, contract_id: str) -> Optional[str]:
        """
        Экспортирует договор в Word формат
        
        Returns:
            Путь к файлу или None если не удалось
        """
        if not HAS_DOCX:
            return None
        
        contract_text = self.get_contract(contract_id)
        if not contract_text:
            return None
        
        metadata = self._load_metadata()
        if contract_id not in metadata:
            return None
        
        # Создаем документ Word
        doc = Document()
        
        # Добавляем заголовок
        title = doc.add_heading(
            metadata[contract_id]['contract_type'],
            level=1
        )
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Добавляем дату
        date_para = doc.add_paragraph()
        date_para.add_run(f"Дата: {metadata[contract_id]['created_at'][:10]}").italic = True
        
        # Добавляем разделитель
        doc.add_paragraph()
        
        # Добавляем текст договора
        # Разбиваем на абзацы
        for line in contract_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        
        # Сохраняем файл
        word_filename = f"contract_{contract_id}.docx"
        word_filepath = os.path.join(self.storage_dir, word_filename)
        doc.save(word_filepath)
        
        # Обновляем метаданные
        metadata[contract_id]['formats'].append('docx')
        self._save_metadata(metadata)
        
        return word_filepath
    
    def delete_contract(self, contract_id: str) -> bool:
        """Удаляет договор"""
        metadata = self._load_metadata()
        if contract_id not in metadata:
            return False
        
        # Удаляем файлы
        for fmt in metadata[contract_id]['formats']:
            if fmt == 'txt':
                filename = metadata[contract_id]['filename']
            elif fmt == 'docx':
                filename = f"contract_{contract_id}.docx"
            else:
                continue
            
            filepath = os.path.join(self.storage_dir, filename)
            if os.path.exists(filepath):
                os.remove(filepath)
        
        # Удаляем из метаданных
        del metadata[contract_id]
        self._save_metadata(metadata)
        
        return True
    
    def get_statistics(self, user_id: Optional[int] = None) -> Dict:
        """Получает статистику по договорам"""
        metadata = self._load_metadata()
        
        stats = {
            'total_contracts': len(metadata),
            'by_type': {},
            'by_user': {}
        }
        
        for contract_id, info in metadata.items():
            # По типам
            contract_type = info['contract_type']
            stats['by_type'][contract_type] = stats['by_type'].get(contract_type, 0) + 1
            
            # По пользователям
            uid = info['user_id']
            if uid not in stats['by_user']:
                stats['by_user'][uid] = {
                    'username': info['username'],
                    'count': 0
                }
            stats['by_user'][uid]['count'] += 1
        
        if user_id:
            return stats['by_user'].get(user_id, {'count': 0})
        
        return stats


# Пример использования
if __name__ == "__main__":
    manager = ContractManager()
    
    # Пример сохранения
    contract_text = "ДОГОВОР\n\nСтороны: ...\n\nПредмет договора: ..."
    contract_id = manager.save_contract(
        contract_text,
        "ФИО: Иван Петров\nДолжность: Директор",
        "Договор купли-продажи",
        123456789,
        "Ivan"
    )
    print(f"Договор сохранён с ID: {contract_id}")
    
    # Пример получения
    retrieved = manager.get_contract(contract_id)
    print(f"Договор найден: {bool(retrieved)}")
    
    # Пример экспорта в Word
    if HAS_DOCX:
        word_path = manager.export_to_word(contract_id)
        print(f"Word файл создан: {word_path}")
